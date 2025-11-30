# src/report_utils.py
import io
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from docx import Document
from docx.shared import Inches
from scipy import stats
import warnings
warnings.filterwarnings('ignore')

# ==== Advanced Visualization Functions ====

def create_distribution_plot(df, column, plot_type='histogram'):
    """
    Create distribution plots for numeric columns with improved error handling
    """
    try:
        print(f"DEBUG: Starting plot creation for column '{column}' with plot_type '{plot_type}'")
        
        # Check if column exists
        if column not in df.columns:
            print(f"ERROR: Column '{column}' not found in dataframe")
            return None
            
        # Get the column data
        col_data = df[column]
        print(f"DEBUG: Column dtype: {col_data.dtype}")
        
        # Enhanced numeric type detection
        is_numeric = False
        if pd.api.types.is_numeric_dtype(col_data):
            is_numeric = True
        elif col_data.dtype in ['int8', 'int16', 'int32', 'int64', 'float16', 'float32', 'float64']:
            is_numeric = True
        elif col_data.dtype == 'object':
            try:
                numeric_converted = pd.to_numeric(col_data, errors='coerce')
                non_null_converted = numeric_converted.dropna()
                if len(non_null_converted) > 0:
                    col_data = numeric_converted
                    is_numeric = True
                    print(f"DEBUG: Column converted from object to numeric")
            except:
                pass
        
        if not is_numeric:
            print(f"ERROR: Column '{column}' is not numeric")
            return None
        
        # Clean data
        plot_data = col_data.dropna()
        if len(plot_data) < 2:
            print(f"ERROR: Insufficient data points ({len(plot_data)})")
            return None
        
        # Create plot based on type
        if plot_type == 'histogram':
            unique_vals = plot_data.nunique()
            nbins = min(50, max(10, int(np.sqrt(len(plot_data)))))
            
            fig = px.histogram(
                x=plot_data, 
                title=f"Distribution of {column}", 
                nbins=nbins,
                labels={'x': column, 'count': 'Frequency'}
            )
        elif plot_type == 'box':
            fig = px.box(y=plot_data, title=f"Box Plot of {column}")
            fig.update_layout(yaxis_title=column)
        elif plot_type == 'violin':
            fig = px.violin(y=plot_data, title=f"Violin Plot of {column}")
            fig.update_layout(yaxis_title=column)
        else:
            # Default to histogram
            nbins = min(50, max(10, int(np.sqrt(len(plot_data)))))
            fig = px.histogram(x=plot_data, title=f"Distribution of {column}", nbins=nbins)
        
        fig.update_layout(showlegend=False, height=400, title_x=0.5, template="plotly_white")
        return fig
        
    except Exception as e:
        print(f"ERROR: Failed to create plot: {e}")
        return None

def create_categorical_analysis(df, column, top_n=10):
    """
    Create comprehensive categorical analysis plots
    """
    try:
        # Check if column exists
        if column not in df.columns:
            print(f"Column '{column}' not found in dataframe")
            return None
            
        # Check if column has categorical data
        if not pd.api.types.is_categorical_dtype(df[column]) and not pd.api.types.is_object_dtype(df[column]):
            print(f"Column '{column}' is not categorical (dtype: {df[column].dtype})")
            return None
        
        # Check if column has any non-null values
        if df[column].isna().all():
            print(f"Column '{column}' has all null values")
            return None
        
        # Remove null values and get value counts
        plot_data = df[column].dropna()
        if len(plot_data) == 0:
            print(f"Column '{column}' has no valid data after removing nulls")
            return None
            
        value_counts = plot_data.value_counts().head(top_n)
        
        if len(value_counts) == 0:
            print(f"Column '{column}' has no value counts after filtering")
            return None
        
        # Check if we have enough data points
        if len(value_counts) < 1:
            print(f"Column '{column}' has insufficient categories ({len(value_counts)})")
            return None
        
        # Create subplot with pie chart and bar chart
        fig = make_subplots(
            rows=1, cols=2,
            specs=[[{"type": "pie"}, {"type": "bar"}]],
            subplot_titles=(f"Distribution of {column}", f"Top {min(top_n, len(value_counts))} values in {column}")
        )
        
        # Pie chart
        fig.add_trace(
            go.Pie(labels=value_counts.index, values=value_counts.values, name=column),
            row=1, col=1
        )
        
        # Bar chart
        fig.add_trace(
            go.Bar(x=value_counts.index, y=value_counts.values, name=column),
            row=1, col=2
        )
        
        fig.update_layout(height=400, showlegend=False)
        return fig
    except Exception as e:
        print(f"Error creating categorical plot for {column}: {e}")
        return None

def create_correlation_analysis(df):
    """
    Create advanced correlation analysis with multiple visualizations
    """
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    if len(numeric_cols) < 2:
        return None
    
    corr_matrix = df[numeric_cols].corr()
    
    # Create subplots
    fig = make_subplots(
        rows=2, cols=2,
        subplot_titles=("Correlation Heatmap", "Correlation Network", 
                       "Pairwise Scatter Matrix", "Correlation Strength"),
        specs=[[{"type": "heatmap"}, {"type": "scatter"}],
               [{"type": "scatter"}, {"type": "bar"}]]
    )
    
    # Heatmap
    fig.add_trace(
        go.Heatmap(z=corr_matrix.values, 
                  x=corr_matrix.columns, 
                  y=corr_matrix.columns,
                  colorscale='RdBu',
                  zmid=0),
        row=1, col=1
    )
    
    # Get strong correlations for network plot
    strong_corr = []
    for i in range(len(corr_matrix.columns)):
        for j in range(i+1, len(corr_matrix.columns)):
            corr_val = corr_matrix.iloc[i, j]
            if abs(corr_val) > 0.5:  # Strong correlation threshold
                strong_corr.append({
                    'x': corr_matrix.columns[i],
                    'y': corr_matrix.columns[j],
                    'correlation': corr_val
                })
    
    if strong_corr:
        corr_df = pd.DataFrame(strong_corr)
        fig.add_trace(
            go.Scatter(x=corr_df['x'], y=corr_df['y'], 
                      mode='markers+text',
                      text=[f"{c:.2f}" for c in corr_df['correlation']],
                      textposition="middle center",
                      marker=dict(size=abs(corr_df['correlation']) * 20)),
            row=1, col=2
        )
    
    # Correlation strength bar chart
    corr_pairs = []
    for i in range(len(corr_matrix.columns)):
        for j in range(i+1, len(corr_matrix.columns)):
            corr_val = corr_matrix.iloc[i, j]
            corr_pairs.append({
                'pair': f"{corr_matrix.columns[i]} - {corr_matrix.columns[j]}",
                'correlation': abs(corr_val)
            })
    
    corr_pairs_df = pd.DataFrame(corr_pairs).sort_values('correlation', ascending=True)
    fig.add_trace(
        go.Bar(x=corr_pairs_df['correlation'], y=corr_pairs_df['pair'], 
               orientation='h', name='Correlation Strength'),
        row=2, col=2
    )
    
    fig.update_layout(height=800, showlegend=False)
    return fig

def create_bivariate_analysis(df, col1, col2):
    """
    Create bivariate analysis plots
    """
    if pd.api.types.is_numeric_dtype(df[col1]) and pd.api.types.is_numeric_dtype(df[col2]):
        # Numeric vs Numeric
        fig = px.scatter(df, x=col1, y=col2, title=f"{col1} vs {col2}",
                        trendline="ols", marginal_x="histogram", marginal_y="histogram")
        return fig
    
    elif df[col1].dtype in ['object', 'category'] and df[col2].dtype in ['object', 'category']:
        # Categorical vs Categorical
        contingency = pd.crosstab(df[col1], df[col2])
        fig = px.imshow(contingency, title=f"Contingency Table: {col1} vs {col2}")
        return fig
    
    else:
        # Mixed types
        if pd.api.types.is_numeric_dtype(df[col1]):
            numeric_col, categorical_col = col1, col2
        else:
            numeric_col, categorical_col = col2, col1
        
        fig = px.box(df, x=categorical_col, y=numeric_col, 
                    title=f"{numeric_col} by {categorical_col}")
        return fig

def create_multivariate_analysis(df):
    """
    Create multivariate analysis plots
    """
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    categorical_cols = df.select_dtypes(include=['object', 'category']).columns
    
    if len(numeric_cols) < 2:
        return None
    
    # Create parallel coordinates plot
    if len(numeric_cols) >= 3:
        fig = px.parallel_coordinates(df[numeric_cols], 
                                    title="Parallel Coordinates Plot")
        return fig
    else:
        # Create scatter matrix
        fig = px.scatter_matrix(df[numeric_cols], 
                               title="Scatter Matrix of Numeric Variables")
        return fig

def create_outlier_analysis(df, column):
    """
    Create outlier analysis plots
    """
    if not pd.api.types.is_numeric_dtype(df[column]):
        return None
    
    # Calculate outliers using IQR method
    Q1 = df[column].quantile(0.25)
    Q3 = df[column].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    
    outliers = df[(df[column] < lower_bound) | (df[column] > upper_bound)]
    
    fig = make_subplots(
        rows=1, cols=2,
        subplot_titles=(f"Box Plot with Outliers", f"Outlier Analysis for {column}")
    )
    
    # Box plot
    fig.add_trace(
        go.Box(y=df[column], name=column, boxpoints='outliers'),
        row=1, col=1
    )
    
    # Outlier scatter plot
    fig.add_trace(
        go.Scatter(x=list(range(len(outliers))), y=outliers[column], 
                  mode='markers', name='Outliers', marker=dict(color='red')),
        row=1, col=2
    )
    
    fig.update_layout(height=400, showlegend=False)
    return fig

# ==== Legacy Functions for Backward Compatibility ====

def top_category_bar(df: pd.DataFrame, col: str, top_n: int = 10):
    """
    Returns a Plotly bar figure showing top values for a categorical column.
    """
    top = df[col].value_counts().nlargest(top_n).reset_index()
    top.columns = [col, 'count']
    fig = px.bar(top, x=col, y='count', title=f"Top {top_n} values in {col}")
    return fig

def time_series_if_exists(df: pd.DataFrame, date_col: str, value_col: str):
    """
    Returns a Plotly line figure if date_col parses as datetime and numeric values exist.
    """
    df_copy = df.copy()
    df_copy[date_col] = pd.to_datetime(df_copy[date_col], errors='coerce')
    ts = df_copy.dropna(subset=[date_col, value_col]).sort_values(date_col)
    fig = px.line(ts, x=date_col, y=value_col, title=f"{value_col} over time")
    return fig

def correlation_heatmap(df: pd.DataFrame):
    """
    Creates a matplotlib heatmap of numeric correlations and returns a BytesIO PNG.
    Returns None if there are <2 numeric columns.
    """
    numeric = df.select_dtypes(include='number')
    if numeric.shape[1] < 2:
        return None

    corr = numeric.corr()

    # dynamic size: avoid too-small or too-large images
    n = corr.shape[0]
    width = max(6, min(12, 0.6 * n))
    height = max(4, min(10, 0.6 * n))

    fig, ax = plt.subplots(figsize=(width, height))
    sns.heatmap(corr, annot=True, fmt=".2f", cmap='coolwarm', ax=ax, square=False,
                cbar_kws={'shrink': 0.7})
    plt.title("Correlation matrix")
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150)
    buf.seek(0)
    plt.close(fig)
    return buf

# ==== Simple fallback insight (keeps app working when LLM unavailable) ====

def build_dataset_stats(df: pd.DataFrame) -> dict:
    """
    Build a dictionary of dataset statistics for summaries and insights.
    This avoids zeros by actually computing metadata.
    """
    stats = {}
    stats["n_rows"] = len(df)
    stats["n_cols"] = len(df.columns)

    # Missing values per column
    stats["missing_per_col"] = df.isna().sum().to_dict()

    # Top values per categorical column (limit to top 3)
    cat_cols = df.select_dtypes(include=["object", "category"]).columns
    cat_top_values = {}
    for col in cat_cols:
        top_vals = df[col].value_counts().head(3).to_dict()
        cat_top_values[col] = top_vals
    stats["categorical_top_values"] = cat_top_values

    # Correlation matrix (safe even if no numeric columns exist)
    num_df = df.select_dtypes(include="number")
    stats["correlation"] = num_df.corr() if not num_df.empty else pd.DataFrame()

    return stats



def simple_template_insight(stats: dict):
    """Return a plain-English one-paragraph summary from stats dict."""
    s = f"The dataset has {stats.get('n_rows', '?')} rows and {stats.get('n_cols', '?')} columns. "
    missing = {k: v for k, v in stats.get('missing_per_col', {}).items() if v > 0}
    if missing:
        s += "Columns with missing values: " + ", ".join([f"{k}({v})" for k, v in missing.items()]) + ". "
    top_cats = stats.get('categorical_top_values', {})
    if top_cats:
        keys = list(top_cats.keys())[:3]
        for c in keys:
            vals = ", ".join([f"{k}({v})" for k, v in top_cats[c].items()])
            s += f"Top values for {c}: {vals}. "
    return s

# ==== Report generation ====

def create_target_based_groupby_analysis(df, target_variable, feature_variable):
    """
    Create comprehensive groupby analysis comparing a feature with the target variable
    """
    try:
        if target_variable not in df.columns or feature_variable not in df.columns:
            return None
        
        # Determine analysis type based on data types
        target_is_numeric = pd.api.types.is_numeric_dtype(df[target_variable])
        feature_is_numeric = pd.api.types.is_numeric_dtype(df[feature_variable])
        
        if target_is_numeric and not feature_is_numeric:
            # Numeric target vs Categorical feature
            groupby_stats = df.groupby(feature_variable)[target_variable].agg([
                'count', 'mean', 'median', 'std', 'min', 'max'
            ]).round(3)
            
            # Create visualization
            fig = px.box(df, x=feature_variable, y=target_variable, 
                        title=f"{target_variable} by {feature_variable}")
            
            return {
                'type': 'numeric_target_categorical_feature',
                'stats': groupby_stats,
                'visualization': fig,
                'target_variable': target_variable,
                'feature_variable': feature_variable
            }
            
        elif not target_is_numeric and feature_is_numeric:
            # Categorical target vs Numeric feature
            groupby_stats = df.groupby(target_variable)[feature_variable].agg([
                'count', 'mean', 'median', 'std', 'min', 'max'
            ]).round(3)
            
            # Create visualization
            fig = px.box(df, x=target_variable, y=feature_variable, 
                        title=f"{feature_variable} by {target_variable}")
            
            return {
                'type': 'categorical_target_numeric_feature',
                'stats': groupby_stats,
                'visualization': fig,
                'target_variable': target_variable,
                'feature_variable': feature_variable
            }
            
        elif not target_is_numeric and not feature_is_numeric:
            # Categorical target vs Categorical feature
            contingency_table = pd.crosstab(df[feature_variable], df[target_variable])
            
            # Create visualization
            fig = px.imshow(contingency_table, 
                           title=f"Contingency Table: {feature_variable} vs {target_variable}",
                           text_auto=True)
            
            return {
                'type': 'categorical_target_categorical_feature',
                'stats': contingency_table,
                'visualization': fig,
                'target_variable': target_variable,
                'feature_variable': feature_variable
            }
            
        else:
            # Numeric target vs Numeric feature
            corr_coef = df[target_variable].corr(df[feature_variable])
            
            # Create visualization
            fig = px.scatter(df, x=feature_variable, y=target_variable, 
                           title=f"{target_variable} vs {feature_variable}",
                           trendline="ols")
            
            return {
                'type': 'numeric_target_numeric_feature',
                'stats': {'correlation': corr_coef},
                'visualization': fig,
                'target_variable': target_variable,
                'feature_variable': feature_variable
            }
            
    except Exception as e:
        print(f"Error creating target-based groupby analysis: {e}")
        return None

def generate_target_comparison_insights(df, target_variable, feature_variable, analysis_result=None):
    """
    Generate 2-pointer insights comparing how the feature affects the target variable
    """
    try:
        insights = []
        
        if analysis_result is None:
            analysis_result = create_target_based_groupby_analysis(df, target_variable, feature_variable)
        
        if analysis_result is None:
            return "Unable to generate insights for this feature-target combination."
        
        analysis_type = analysis_result['type']
        target_var = analysis_result['target_variable']
        feature_var = analysis_result['feature_variable']
        
        insights.append(f"🎯 **Target Variable Impact Analysis: {feature_var} → {target_var}**")
        insights.append("")
        
        if analysis_type == 'numeric_target_categorical_feature':
            stats = analysis_result['stats']
            
            # Pointer 1: Statistical Impact
            mean_values = stats['mean']
            std_values = stats['std']
            
            best_category = mean_values.idxmax()
            worst_category = mean_values.idxmin()
            best_mean = mean_values[best_category]
            worst_mean = mean_values[worst_category]
            impact_range = best_mean - worst_mean
            
            insights.append("📊 **Pointer 1: Statistical Impact on Target Variable**")
            insights.append(f"• **{feature_var}** shows significant impact on **{target_var}**")
            insights.append(f"• **Best performing category**: {best_category} (mean: {best_mean:.2f})")
            insights.append(f"• **Worst performing category**: {worst_category} (mean: {worst_mean:.2f})")
            insights.append(f"• **Impact range**: {impact_range:.2f} units difference")
            
            if impact_range > std_values.mean():
                insights.append(f"• **High impact**: Range exceeds average variability ({std_values.mean():.2f})")
            else:
                insights.append(f"• **Moderate impact**: Range within normal variability")
            
            # Pointer 2: Business/Strategic Implications
            insights.append("")
            insights.append("💡 **Pointer 2: Strategic Implications for Target Variable**")
            
            # Calculate coefficient of variation for stability assessment
            cv_values = (std_values / mean_values).abs()
            most_stable = cv_values.idxmin()
            least_stable = cv_values.idxmax()
            
            insights.append(f"• **Most stable category**: {most_stable} (lowest variability)")
            insights.append(f"• **Least stable category**: {least_stable} (highest variability)")
            
            # Sample size considerations
            count_values = stats['count']
            min_samples = count_values.min()
            max_samples = count_values.max()
            
            if min_samples < 30:
                insights.append(f"• **⚠️ Warning**: Some categories have small sample sizes (min: {min_samples})")
            else:
                insights.append(f"• **✅ Good sample sizes**: All categories have adequate data (min: {min_samples})")
            
            # Actionable insights
            if impact_range > 2 * std_values.mean():
                insights.append(f"• **🎯 Action**: Focus on improving {worst_category} category performance")
                insights.append(f"• **📈 Opportunity**: {best_category} category shows best practices to replicate")
            else:
                insights.append(f"• **📊 Observation**: {feature_var} has moderate influence on {target_var}")
                insights.append(f"• **🔍 Next step**: Investigate other factors for stronger impact")
        
        elif analysis_type == 'categorical_target_numeric_feature':
            stats = analysis_result['stats']
            
            # Pointer 1: Feature Distribution by Target
            insights.append("📊 **Pointer 1: Feature Distribution Impact on Target Variable**")
            
            target_categories = stats.index.tolist()
            feature_means = stats['mean']
            feature_stds = stats['std']
            
            # Find which target categories have significantly different feature values
            max_mean = feature_means.max()
            min_mean = feature_means.min()
            range_diff = max_mean - min_mean
            
            best_target = feature_means.idxmax()
            worst_target = feature_means.idxmin()
            
            insights.append(f"• **{feature_var}** varies significantly across **{target_var}** categories")
            insights.append(f"• **Highest feature values**: {best_target} (mean: {max_mean:.2f})")
            insights.append(f"• **Lowest feature values**: {worst_target} (mean: {min_mean:.2f})")
            insights.append(f"• **Feature range**: {range_diff:.2f} units difference")
            
            # Pointer 2: Predictive Power and Thresholds
            insights.append("")
            insights.append("💡 **Pointer 2: Predictive Power and Decision Thresholds**")
            
            # Calculate separation between categories
            if len(target_categories) == 2:
                # Binary classification case
                cat1, cat2 = target_categories
                mean1, mean2 = feature_means[cat1], feature_means[cat2]
                std1, std2 = feature_stds[cat1], feature_stds[cat2]
                
                # Calculate separation (Cohen's d)
                pooled_std = np.sqrt((std1**2 + std2**2) / 2)
                cohens_d = abs(mean1 - mean2) / pooled_std
                
                insights.append(f"• **Separation strength**: {cohens_d:.2f} (Cohen's d)")
                if cohens_d > 0.8:
                    insights.append("• **Strong separation**: Feature is highly predictive")
                elif cohens_d > 0.5:
                    insights.append("• **Moderate separation**: Feature has good predictive power")
                else:
                    insights.append("• **Weak separation**: Feature has limited predictive power")
                
                # Suggest threshold
                threshold = (mean1 + mean2) / 2
                insights.append(f"• **Suggested threshold**: {threshold:.2f} for classification")
            else:
                # Multi-class case
                insights.append(f"• **Multi-class analysis**: {len(target_categories)} target categories")
                insights.append(f"• **Feature discrimination**: {range_diff:.2f} range across categories")
                
                if range_diff > 2 * feature_stds.mean():
                    insights.append("• **High discrimination**: Feature strongly differentiates target categories")
                else:
                    insights.append("• **Moderate discrimination**: Feature provides some differentiation")
        
        elif analysis_type == 'categorical_target_categorical_feature':
            contingency = analysis_result['stats']
            
            # Pointer 1: Association Strength
            insights.append("📊 **Pointer 1: Association Strength Between Variables**")
            
            # Calculate chi-square test
            from scipy.stats import chi2_contingency
            chi2, p_value, dof, expected = chi2_contingency(contingency)
            
            insights.append(f"• **Chi-square statistic**: {chi2:.3f}")
            insights.append(f"• **P-value**: {p_value:.3f}")
            insights.append(f"• **Degrees of freedom**: {dof}")
            
            if p_value < 0.001:
                insights.append("• **Very strong association**: Highly significant relationship")
            elif p_value < 0.05:
                insights.append("• **Strong association**: Statistically significant relationship")
            else:
                insights.append("• **Weak association**: No significant relationship")
            
            # Pointer 2: Practical Implications
            insights.append("")
            insights.append("💡 **Pointer 2: Practical Implications for Target Variable**")
            
            # Find strongest associations
            total = contingency.sum().sum()
            expected_prop = expected / total
            observed_prop = contingency / total
            
            # Calculate standardized residuals
            residuals = (observed_prop - expected_prop) / np.sqrt(expected_prop)
            
            # Find most over-represented combinations
            max_residual = residuals.max().max()
            min_residual = residuals.min().min()
            
            insights.append(f"• **Association strength range**: {min_residual:.2f} to {max_residual:.2f} (standardized residuals)")
            
            if abs(max_residual) > 2:
                insights.append("• **Strong patterns**: Some feature-target combinations are highly over-represented")
            elif abs(max_residual) > 1:
                insights.append("• **Moderate patterns**: Some feature-target combinations show notable associations")
            else:
                insights.append("• **Weak patterns**: Feature-target combinations are mostly random")
            
            # Sample size considerations
            min_count = contingency.min().min()
            if min_count < 5:
                insights.append(f"• **⚠️ Warning**: Some cells have small counts (min: {min_count}) - results may be unreliable")
            else:
                insights.append(f"• **✅ Good sample sizes**: All combinations have adequate data")
        
        else:  # numeric_target_numeric_feature
            corr_coef = analysis_result['stats']['correlation']
            
            # Pointer 1: Correlation Strength and Direction
            insights.append("📊 **Pointer 1: Linear Relationship Strength**")
            insights.append(f"• **Correlation coefficient**: {corr_coef:.3f}")
            
            if abs(corr_coef) > 0.7:
                strength = "very strong"
            elif abs(corr_coef) > 0.5:
                strength = "strong"
            elif abs(corr_coef) > 0.3:
                strength = "moderate"
            else:
                strength = "weak"
            
            direction = "positive" if corr_coef > 0 else "negative"
            insights.append(f"• **Relationship**: {strength} {direction} linear relationship")
            
            if abs(corr_coef) > 0.7:
                insights.append("• **High predictive power**: Feature strongly predicts target variable")
            elif abs(corr_coef) > 0.3:
                insights.append("• **Moderate predictive power**: Feature has some predictive value")
            else:
                insights.append("• **Low predictive power**: Feature has limited linear relationship")
            
            # Pointer 2: Practical Implications and Next Steps
            insights.append("")
            insights.append("💡 **Pointer 2: Strategic Implications and Next Steps**")
            
            if abs(corr_coef) > 0.5:
                insights.append(f"• **🎯 Action**: {feature_var} is a key driver of {target_variable}")
                if corr_coef > 0:
                    insights.append(f"• **📈 Strategy**: Increase {feature_var} to improve {target_variable}")
                else:
                    insights.append(f"• **📉 Strategy**: Decrease {feature_var} to improve {target_variable}")
            else:
                insights.append(f"• **🔍 Investigation**: {feature_var} has limited direct impact on {target_variable}")
                insights.append(f"• **📊 Next step**: Consider non-linear relationships or interaction effects")
            
            # R-squared interpretation
            r_squared = corr_coef ** 2
            insights.append(f"• **Explained variance**: {r_squared:.1%} of {target_variable} variance explained by {feature_var}")
            
            if r_squared > 0.5:
                insights.append("• **High explanatory power**: Feature accounts for majority of target variation")
            elif r_squared > 0.25:
                insights.append("• **Moderate explanatory power**: Feature explains significant portion of target variation")
            else:
                insights.append("• **Low explanatory power**: Feature explains limited target variation")
        
        return "\n".join(insights)
        
    except Exception as e:
        return f"Error generating comparison insights: {e}"

def create_word_report(insights_text: str, img_buffers: dict, out_path: str = "report/sample_report.docx"):
    """
    Create a simple Word document with the insights paragraph and supplied image buffers.
    img_buffers: dict where key=title, value=BytesIO (seeked to 0)
    """
    doc = Document()
    doc.add_heading('Executive Summary — AI Data Storyteller', level=1)
    doc.add_paragraph(insights_text)
    doc.add_heading('Visualizations', level=2)

    for title, buf in img_buffers.items():
        doc.add_paragraph(title)
        buf.seek(0)
        try:
            doc.add_picture(buf, width=Inches(6))
        except Exception:
            # fallback: save to temp and insert
            tmp_path = "temp_vis.png"
            with open(tmp_path, "wb") as f:
                f.write(buf.getbuffer())
            doc.add_picture(tmp_path, width=Inches(6))
    doc.save(out_path)
    return out_path
